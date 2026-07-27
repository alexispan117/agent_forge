"""
生成美化版简历 — 基于 AgentForge 真实项目状态
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT = Path(r"D:\hermes\work\agentgroup\ai-knowledge")
OUT.mkdir(parents=True, exist_ok=True)
doc = Document()

# ── 页面设置 ──
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ── 颜色 ──
PRIMARY = RGBColor(0x63, 0x66, 0xF1)    # Indigo
DARK = RGBColor(0x1E, 0x29, 0x3B)       # Slate 800
MUTED = RGBColor(0x64, 0x74, 0x8B)      # Slate 500
ACCENT = RGBColor(0x10, 0xB9, 0x81)     # Emerald
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BORDER = RGBColor(0xE2, 0xE8, 0xF0)

# ── 样式 ──
style = doc.styles["Normal"]
style.font.name = "Inter"
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.35
style.font.color.rgb = DARK

for lv in range(1, 4):
    s = doc.styles[f"Heading {lv}"]
    s.font.name = "Inter"
    s.font.color.rgb = PRIMARY

# ── 辅助函数 ──
def add_section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY
    run.font.name = "Inter"
    # 下划线
    p_border = p._element.get_or_add_pPr()
    pBdr = p_border.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '6',
        qn('w:space'): '4',
        qn('w:color'): '6366F1',
    })
    pBdr.append(bottom)
    p_border.append(pBdr)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("▸ ")
    run.font.color.rgb = PRIMARY
    run.font.size = Pt(10)
    if bold_prefix:
        run2 = p.add_run(bold_prefix)
        run2.font.bold = True
        run2.font.size = Pt(10)
        run2.font.color.rgb = DARK
        run3 = p.add_run(text)
        run3.font.size = Pt(10)
        run3.font.color.rgb = DARK
    else:
        run2 = p.add_run(text)
        run2.font.size = Pt(10)
        run2.font.color.rgb = DARK
    return p

def add_meta_line(items):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    for i, item in enumerate(items):
        if i > 0:
            sep = p.add_run("  |  ")
            sep.font.size = Pt(10)
            sep.font.color.rgb = MUTED
        run = p.add_run(item)
        run.font.size = Pt(10)
        run.font.color.rgb = MUTED
    return p

# ═══════════ 正文 ═══════════

# ── 姓名 + 标题 ──
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(2)
run = title.add_run("Edward (张明远)")
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = DARK

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(6)
run = subtitle.add_run("AI Agent 应用开发 / AI 产品经理")
run.font.size = Pt(13)
run.font.color.rgb = PRIMARY
run.font.bold = True

add_meta_line([
    "138-xxxx-xxxx",
    "zhangmingyuan@email.com",
    "上海",
    "github.com/agentforge",
    "2003.07",
])

# ── 专业概述 ──
add_section_heading("专业概述")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.3)
run = p.add_run(
    "信息工程专业应届毕业生，独立完成企业级多智能体编排平台 AgentForge 从 0 到 1 的全栈建设。"
    "项目以金融/制造/政务合规审计为场景，实现 Supervisor-Worker 四节点微服务架构、MCP/A2A 双协议 Agent 互联、"
    "CLEAR 五维评估体系与全链路异步化。具备 30 分钟企业路演能力——"
    "能现场演示复杂任务原子化拆解、故障注入自愈（杀死 Worker 容器后业务不中断）、"
    "以及离线 Mock 零网络演示。技术栈覆盖 async Python/FastAPI 后端、LangGraph 编排、"
    "Docker-Compose 容器编排、Jinja2 + ECharts 仪表盘。"
)
run.font.size = Pt(10)
run.font.color.rgb = DARK

# ── 能力矩阵 ──
add_section_heading("能力矩阵")
doc.add_paragraph()

# 表格：3 列
table = doc.add_table(rows=3, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

headers = ["🏗️ 架构与工程", "🤖 AI / Agent", "🛠️ 工具与平台"]
data = [
    ["FastAPI + asyncio 异步后端\nDocker-Compose 微服务编排\nSQLAlchemy + ChromaDB 持久化\nGitHub Actions CI + pytest", "Prompt Engineering (CoT/ReAct)\nRAG 混合检索 + RRF 融合\nLangGraph 状态机编排\nCLEAR 五维评估体系", "Python / C++\nLangChain / LangGraph\nOpenAI / DeepSeek API\nDocker / Git / Linux"],
    ["MCP / A2A Agent 协议", "@with_fallback 故障自愈", "ECharts + SSE 实时仪表盘"],
]

for i, row in enumerate(table.rows):
    for j, cell in enumerate(row.cells):
        if i == 0:
            p = cell.paragraphs[0]
            p.add_run(headers[j]).font.bold = True
        else:
            cell.text = data[i-1][j] if i-1 < len(data) and j < len(data[i-1]) else ""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = DARK

# Set column widths
for row in table.rows:
    row.cells[0].width = Cm(5.5)
    row.cells[1].width = Cm(5.5)
    row.cells[2].width = Cm(5.5)

# ── 核心项目 ──
add_section_heading("核心项目")

# AgentForge
p = doc.add_paragraph()
run = p.add_run("AgentForge — 企业级多智能体编排平台")
run.font.bold = True
run.font.size = Pt(12)
run.font.color.rgb = DARK

add_meta_line(["2026.03 – 2026.07", "独立架构设计 + 全栈开发", "可 30 分钟路演的商业化产品"])

bullets = [
    ("架构与工程：",
     "Supervisor + 3 Worker 四容器微服务架构（Docker-Compose），代码卷挂载支持热重载。"
     "自研 MCP (JSON-RPC) + A2A (AgentCard 服务发现) 双协议层，Agent 间零硬编码互联。"
     "LangGraph 状态机驱动 Supervisor 意图分类与 DAG 拆解。"
     "asyncio.gather 实现 Worker 真并行调度。"),
    ("韧性体系：",
     "@with_fallback 三级防线——指数退避重试 → Circuit Breaker 熔断 → 本地降级兜底。"
     "演示中杀死分析 Worker 容器后，工作流 100% 交付（降级路径），重启后自动恢复满分。"
     "各 I/O 调用标注 # ASYNC 注释并锁定库版本。"),
    ("评估与质量：",
     "CLEAR 五维评估（成本/延迟/效能/保证/可靠性），全部由真实执行数据计算，ECharts 雷达图实时渲染。"
     "21 个 pytest 用例 + GitHub Actions CI（全模块断链冒烟）。"
     "密钥全量外移至 .env 环境变量，config.yaml 仅存 ${VAR} 占位符。"),
    ("演示友好设计：",
     "离线 Mock 模式（LLM_MOCK_MODE=true）默认开启，6 个场景的预置响应模板覆盖拆解/异常/脱敏/报告等全流程，"
     "零网络也能完美演示。demo.sh 一键启动 + 预热种子数据 + 自动打开浏览器。"
     "前端 Jinja2 + SSE 实时推送 + ECharts 白底 SaaS 仪表盘（Inter 字体本地加载）。"),
    ("量化成果：",
     "复杂多步任务完成率从 < 50% 提升至 ~75%（21 个自动化用例验证）"
     "；RAG 增强后垂直领域问答幻觉率降低约 15%；"
     "四容器联调全链路 200 通过；CLEAR 评分如实反映故障注入。"),
]
for bold, text in bullets:
    add_bullet(text, bold)

# 智能小车
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run("智能小车自主决策系统（嵌入式 C++ / 本科科研）")
run.font.bold = True
run.font.size = Pt(11)

add_meta_line(["2024.06 – 2024.12", "独立开发"])
add_bullet("基于多传感器（超声波 / 红外）融合的「感知→决策→执行」闭环。"
           "状态机实现避障 / 循迹 / 路径规划动态切换。"
           "架构与 Agent 的 Planning–Tool Use–Execution 范式同源。")

# 数据库项目
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run("数据库系统设计与优化（课程项目）")
run.font.bold = True
run.font.size = Pt(11)

add_meta_line(["2025.09 – 2025.10"])
add_bullet("SQL Server 架构设计、索引优化与客户端开发。"
           "为后续 ChromaDB 向量检索与混合检索调优打下基础。")

# ── 实习经历 ──
add_section_heading("实习经历")

p = doc.add_paragraph()
run = p.add_run("金工实习 — 智能制造流程实践")
run.font.bold = True
run.font.size = Pt(11)
run.font.color.rgb = DARK
add_meta_line(["上海大学工程训练中心", "2023.06 – 2023.08"])
add_bullet("参与数控机床、3D 打印、CATIA 设计的完整工业制造流程，"
           "理解制造业现场的真实痛点，为「AI + 工业制造」场景落地提供一线认知。")

# ── 教育背景 ──
add_section_heading("教育背景")

p = doc.add_paragraph()
run = p.add_run("上海大学")
run.font.bold = True
run.font.size = Pt(11)
run2 = p.add_run(" · 信息工程专业（本科） · 2022.09 – 2026.06")
run2.font.size = Pt(11)

add_bullet("主修：数据结构与算法、数据库系统、机器学习基础、信号与系统、嵌入式系统设计")
add_bullet("研究方向：大语言模型应用开发与 Agent 架构设计")
add_bullet("英语：CET-6，熟练阅读英文技术文档与论文")

# ── 技术影响力 ──
add_section_heading("技术影响力")

add_bullet("持续追踪并实践：MCP / A2A 协议（已落地）| LangGraph 状态图编排 | AutoGen / CrewAI | Harness Engineering | Reflexion 反思机制")
add_bullet("技术布道：独立生成 20,000+ 字 AI Agent 教学文档（涵盖四项目逐行代码讲解 + 面试题库 + 行业案例）")
add_bullet("开源：AgentForge 完整企业级项目可供 CTO 现场审查源码")

# ── 保存 ──
out_path = OUT / "Edward_Resume_AI_Agent_2026.docx"
doc.save(str(out_path))
print(f"✅ 简历已生成: {out_path}")
print(f"   大小: {out_path.stat().st_size // 1024}KB")
