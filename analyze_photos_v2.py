"""
📸 图像分析报告 v2 — Qwen-VL 视觉提取 + .docx 生成
"""
import re, time, json, base64
from pathlib import Path
from datetime import datetime
from collections import defaultdict

from openai import OpenAI
from PIL import Image
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

PHOTO_DIR = Path(r"D:\hermes\work\agentgroup\ai-knowledge\photo1")
OUT_DIR = Path(r"D:\hermes\work\agentgroup\ai-knowledge")
OUT_DIR.mkdir(parents=True, exist_ok=True)

# DashScope API（阿里百炼）
import yaml
cfg = yaml.safe_load(open(Path(__file__).parent / "config.yaml"))
API_KEY = cfg["embedding"]["api_key"]
client = OpenAI(api_key=API_KEY, base_url="https://dashscope.aliyuncs.com/compatible-mode/v1")
MODEL = "qwen-vl-plus"

exts = {".jpg", ".jpeg", ".png", ".webp", ".bmp"}
images = sorted([p for p in PHOTO_DIR.iterdir() if p.suffix.lower() in exts])

PROMPT = """请详细提取这张图片中的所有文字内容，按原文格式输出。
如果是截图，请完整提取可见的所有文字。
如果包含表格，用表格格式输出。
如果有多页内容，请全部提取。
只输出提取的文字内容，不要添加分析评论。"""


def analyze(img_path: Path) -> dict:
    img = Image.open(img_path)
    w, h = img.size
    size_kb = img_path.stat().st_size / 1024
    b64 = base64.b64encode(open(img_path, "rb").read()).decode()

    # 日期
    m = re.search(r"(\d{4})(\d{2})(\d{2})", img_path.stem)
    date = f"{m.group(1)}-{m.group(2)}-{m.group(3)}" if m else "未知"
    tm = re.search(r"_(\d{6})", img_path.stem)
    time_str = f"{tm.group(1)[:2]}:{tm.group(1)[2:4]}:{tm.group(1)[4:6]}" if tm else ""

    # 视觉提取文字（最多重试 2 次）
    text = ""
    for attempt in range(2):
        try:
            resp = client.chat.completions.create(
                model=MODEL,
                messages=[{"role": "user", "content": [
                    {"type": "text", "text": PROMPT},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}}
                ]}],
                temperature=0.05, max_tokens=2000,
            )
            text = resp.choices[0].message.content.strip()
            break
        except Exception as e:
            text = f"[API 错误] {e}"
            time.sleep(2)

    return {
        "filename": img_path.name, "date": date, "time": time_str,
        "size_kb": round(size_kb, 1), "dim": f"{w}x{h}",
        "cat": "照片" if size_kb > 2000 else "截图",
        "text": text, "text_len": len(text),
    }


# ── 主流程 ──
print(f"共 {len(images)} 张图片，使用 Qwen-VL-Plus 提取文字...\n")

all_results = []
for idx, img in enumerate(images):
    print(f"[{idx+1:2d}/{len(images)}] {img.name[:35]:35s}...", end=" ", flush=True)
    r = analyze(img)
    all_results.append(r)
    preview = r["text"][:60].replace("\n", " ") if r["text"] else "(无文字)"
    print(f"{r['cat']} | {preview}")
    if idx < len(images) - 1:
        time.sleep(0.5)  # 限速

# ── 生成 .docx ──
doc = Document()
s = doc.styles["Normal"]; s.font.name = "微软雅黑"; s.font.size = Pt(11)

# 封面
doc.add_heading("AI 图像分析报告 v2", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f"分析时间: {datetime.now():%Y-%m-%d %H:%M}")
doc.add_paragraph(f"图片来源: {PHOTO_DIR}")
doc.add_paragraph(f"图片总数: {len(images)} 张")
doc.add_paragraph("分析引擎: Qwen-VL-Plus (阿里百炼 DashScope)")
doc.add_page_break()

# 一、概览
doc.add_heading("一、概览", 1)
cats = defaultdict(int)
for r in all_results:
    cats[r["cat"]] += 1
for c, n in sorted(cats.items(), key=lambda x: -x[1]):
    doc.add_paragraph(f"  {c}: {n} 张")

dates = sorted(set(r["date"] for r in all_results if r["date"] != "未知"))
doc.add_paragraph(f"时间跨度: {len(dates)} 天（{dates[0]} ~ {dates[-1]}）" if dates else "日期信息: 无")
total_kb = sum(r["size_kb"] for r in all_results)
doc.add_paragraph(f"总大小: {total_kb:.0f}KB ({total_kb/1024:.1f}MB)")
has_text = sum(1 for r in all_results if r["text"] and not r["text"].startswith("[API"))
doc.add_paragraph(f"成功提取文字: {has_text}/{len(images)} 张")
doc.add_page_break()

# 二、按日期逐张分析
doc.add_heading("二、逐张分析", 1)
for d in dates:
    grp = [r for r in all_results if r["date"] == d]
    doc.add_heading(f"📅 {d}（{len(grp)} 张）", 2)
    for r in grp:
        doc.add_heading(f"  {r['filename']}", 3)
        doc.add_paragraph(f"类型: {r['cat']} | 尺寸: {r['dim']} | 大小: {r['size_kb']}KB")
        if r["text"] and not r["text"].startswith("[API"):
            doc.add_paragraph("提取文字:")
            doc.add_paragraph(r["text"][:3000])
        else:
            doc.add_paragraph(f"（{r['text']}）" if r["text"].startswith("[API") else "（未识别到文字）")
        doc.add_paragraph("")

# 三、文字内容汇总
doc.add_page_break()
doc.add_heading("三、全局文字内容汇总", 1)
all_text_parts = []
for r in all_results:
    if r["text"] and not r["text"].startswith("[API"):
        all_text_parts.append(f"── {r['filename']} ──\n{r['text']}")
combined = "\n\n".join(all_text_parts)
if len(combined) > 100000:
    combined = combined[:100000] + "\n\n…(截断至 100000 字符)"
doc.add_paragraph(combined if combined else "（无可用文字内容）")

# 保存
out_path = OUT_DIR / "图像分析报告_v2.docx"
doc.save(str(out_path))
print(f"\n✅ 报告已保存: {out_path}")
print(f"   共 {len(all_results)} 张 | 成功提取 {has_text} 张")
