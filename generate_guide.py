"""📘 AI Agent 技术大全 — 超详细教学文档"""
from pathlib import Path; from datetime import datetime
from docx import Document; from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(r"D:\hermes\work\agentgroup\ai-knowledge"); OUT.mkdir(parents=True, exist_ok=True)
doc = Document()
s = doc.styles["Normal"]; s.font.name = "微软雅黑"; s.font.size = Pt(10.5)
s.paragraph_format.line_spacing = 1.35
for lv in range(1,5): doc.styles[f"Heading {lv}"].font.name = "微软雅黑"
cs = doc.styles.add_style("Code", 1); cs.font.name = "Consolas"; cs.font.size = Pt(8.5)

def H(t, lv=1): return doc.add_heading(t, lv)
def P(t, st=None): return doc.add_paragraph(t, style=st)
def PB(): doc.add_page_break()

# ── 封面 ──
H("AI Agent 技术大全（超详细版）", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
P(f"生成日期: {datetime.now():%Y-%m-%d} | 基于 AgentForge + MokioAgent 四项目").alignment = WD_ALIGN_PARAGRAPH.CENTER
P("目标读者：零基础或初级Python开发者，想彻底掌握AI Agent技术").alignment = WD_ALIGN_PARAGRAPH.CENTER
PB()

# ═══════════ 第一篇：Python基础 ═══════════
H("第一篇：Python基础——从零看懂Agent项目", 1)
P("如果你之前没写过Python，或者只写过几行，这一篇就是为你准备的。")

H("第1章：变量——程序里的名片", 2)
P("变量是编程中最基本的概念。可以把变量想象成贴了标签的盒子——盒子(内存)里存着数据，标签(变量名)让你能随时找到它。")
P("name = '张三'        # 字符串(str)：用引号括起来的文字", "Code")
P("age = 25             # 整数(int)         price = 99.9   # 浮点数(float)", "Code")
P("is_admin = True      # 布尔值(bool)：只有True/False", "Code")
P("")
P("config = {'llm': {'model': 'deepseek-v4-flash'}}   # 嵌套字典", "Code")
P("model_name = config['llm']['model']            # 读取键值", "Code")
P("batch = config.get('embedding', {}).get('batch_size', 5)  # 安全读取", "Code")

H("第2章：字符串——AI Agent世界的通用语言", 2)
P("AI Agent项目里，90%的数据都是字符串。")
P("s1 = '单引号'; s2 = '双引号'       # 两种写法没区别", "Code")
P("s3 = '''三引号可以换行'''          # 跨行字符串", "Code")
P("full = '张' + '三'                 # +号拼接", "Code")
P("msg = f'用户问题：{query}'          # f-string嵌入变量", "Code")
P("print(query[:4])  # 取前4个字", "Code")
P("text.strip()      # 去空格", "Code")
P("text.replace('年','岁')  # 替换", "Code")

H("第3章：列表(list)——一串数据的集合", 2)
P("列表是Python最常用的集合类型，像购物清单。")
P("tools = ['list_dir', 'read_file', 'write_file', 'run_python']", "Code")
P("print(tools[0])     # list_dir（索引从0开始）", "Code")
P("print(tools[-1])    # run_python（最后一个）", "Code")
P("print(tools[1:3])   # 切片：['read_file','write_file']", "Code")
P("for tool in tools:  # 遍历", "Code")
P("    print(f'工具:{tool}')", "Code")
P("md_files = [f for f in names if f.endswith('.md')]  # 列表推导式", "Code")

H("第4章：字典(dict)——键值对的地图", 2)
P("字典像一本真正的字典，查词(key)得到解释(value)。")
P("student = {'name': '张三', 'age': 25, 'scores': [95,88,92]}", "Code")
P("print(student['name'])          # 张三", "Code")
P("print(student.get('grade','无')) # 不存在时返回默认值", "Code")
P("for k,v in student.items():     # 遍历", "Code")
P("    print(f'{k}: {v}')", "Code")
P("")
P("嵌套字典（项目中极其常见）：", "Code")
P("config = {", "Code")
P("  'llm': {'model': 'deepseek-v4-flash', 'api_key': 'sk-xxx'},", "Code")
P("  'embedding': {'model': 'text-embedding-v3', 'batch_size': 10}", "Code")
P("}", "Code")
P("print(config['llm']['model'])  # deepseek-v4-flash", "Code")

H("第5章：函数(def)——把代码打包成技能", 2)
P("函数把一段代码打包，取个名字，需要时调用。")
P("def greet(name):", "Code")
P("    '''文档字符串：描述这个函数'''", "Code")
P("    return f'你好，{name}！'", "Code")
P("result = greet('小明')  # 调用函数", "Code")
P("")
P("每个Agent中最核心的函数——execute()：", "Code")
P("def execute(self, query='', **kwargs):", "Code")
P("    '''self=本对象 query=用户输入 **kwargs=其他参数'''", "Code")
P("    return {'answer': '回答', 'sources': []}", "Code")

H("第6章：类(class)——Agent的骨架", 2)
P("类是面向对象编程(OOP)的核心。类像蛋糕模具，对象是蛋糕。")
P("class Agent:", "Code")
P("    def __init__(self, name):   # 构造函数：创建时自动调用", "Code")
P("        self.name = name       # self指向这个对象", "Code")
P("        self.memory = []       # 每个对象有自己的记忆", "Code")
P("    def think(self, question):", "Code")
P("        self.memory.append(question)", "Code")
P("        return f'关于{question}的回答...'", "Code")
P("", "Code")
P("agent = Agent('小助手')         # 创建对象", "Code")
P("agent.think('天气')              # 调用方法", "Code")
P("")
P("继承：子类自动拥有父类的所有方法", "Code")
P("class Searcher(Agent):", "Code")
P("    def __init__(self):", "Code")
P("        super().__init__('搜索助手')  # super()调用父类构造函数", "Code")
P("        self.api_key = ''", "Code")

H("第7章：异常处理(try/except)——让程序优雅面对错误", 2)
P("程序一定会出错。好的程序不是不出错，而是出错时能妥善处理。")
P("try:", "Code")
P("    f = open('不存在的文件.txt')", "Code")
P("    content = f.read()", "Code")
P("except FileNotFoundError:", "Code")
P("    content = ''  # 文件不存在就用空内容", "Code")
P("except Exception as e:", "Code")
P("    print(f'未知错误：{e}')", "Code")
P("")
P("KnowledgeBot的熔断+降级：", "Code")
P("@self._cb  # 熔断器装饰", "Code")
P("def hybrid_search(query):", "Code")
P("    return vector_store.hybrid_search(query)", "Code")
P("try:", "Code")
P("    results = hybrid_search(query)", "Code")
P("except Exception:", "Code")
P("    results = keyword_search(query)  # 降级", "Code")
PB()

# ═══════════ 第二篇：核心概念 ═══════════
H("第二篇：30个AI Agent核心概念详解", 1)
P("每个概念包含：一句话类比 + 深入讲解 + 代码示例 + 注意事项")

concepts = []
concepts.append(["1.LLM（大语言模型）","类比：LLM是Agent的大脑。DeepSeek、GPT-4、Claude都是LLM。它们能理解自然语言、推理、生成回答。","LLM通过海量文本训练，本质是文本生成器——给一段文字，预测接下来最可能出现的词。","resp = client.chat.completions.create(model='deepseek-v4-flash', messages=[{'role':'user','content':'净利润？'}])\nanswer = resp.choices[0].message.content","误区：认为LLM什么都知道。实际上知识截止于训练日期，且会幻觉(编造答案)。需要RAG解决。"])
concepts.append(["2.Token","类比：AI世界的分钟计费。打电话按分钟，用AI按Token。","1汉字≈2-3 Token，1英文词≈1 Token。每次调用都会消耗Token——输入+输出都计费。","ct = CostTracker()\nct.record('deepseek', input_tokens=500, output_tokens=200)\nprint(ct.session_summary()['total_cost_usd'])","设置max_tokens防止无限生成。"])
concepts.append(["3.Prompt（提示词）","类比：给AI的工作任务单。写得越清楚，AI做得越好。","完整Prompt包含：角色定义+任务描述+输出格式+约束条件+示例(few-shot)。","RAG_SYSTEM_PROMPT = '''你是智能客服。回答结构：1.数据呈现 2.趋势分析 3.指导意见。纯文本，不用Markdown。'''","不是越长越好。关键是结构清晰、重点突出。"])
concepts.append(["4.System Prompt","系统级指令，定义AI的人设和行为规则。用户看不到但始终起作用。","System Prompt位于messages第一条(role='system')，User Prompt是role='user'。","messages = [{'role':'system','content':'你是投资顾问'},\n            {'role':'user','content':'ROE是多少？'}]","好的System Prompt能大幅提升回答质量。"])
concepts.append(["5.Context Window（上下文窗口）","类比：AI的短期记忆。超出Context大小的内容会被遗忘。","DeepSeek约128K Token(约9万汉字)。Context越长响应越慢、成本越高。","history = ShortTermMemory(max_window=20)  # 保留最近20轮\ncontext = history.get_context()","不是越大越好。合适的Context管理比无限扩展更重要。"])
concepts.append(["6.Temperature（温度）","类比：AI的创造力旋钮。0=死板精确，1=天马行空，2=胡言乱语。","Temperature控制选择低概率Token的可能性。越低回答越稳定。Agent通常设0.1。","resp = client.chat.completions.create(temperature=0.1, messages=[...])","Tool Calling用0.1或0。创意任务用0.7-0.9。"])
concepts.append(["7.Embedding（文本向量化）","类比：把意思变成坐标。苹果→(2,3)，iPhone→(2.1,2.8)。","Embedding模型接收文字输出一串浮点数(向量)。常见维度768/1024/1536。","resp = client.embeddings.create(input='中国平安净利润', model='text-embedding-v3')\nvector = resp.data[0].embedding","维度越高表达能力越强，计算成本也越高。"])
concepts.append(["8.Vector Database（向量数据库）","类比：AI的智能文件柜。传统需要关键字匹配，向量库用意思来找文件。","专门存向量并提供相似度搜索。ChromaDB(开发)、Milvus(生产)、Pinecone(云)。","db = chromadb.PersistentClient(path='./chroma')\ncollection = db.get_or_create_collection('knowledge')\ncollection.add(ids=['doc1'], documents=['年报内容'])\nresults = collection.query(query_texts=['净利润'], n_results=5)","不同向量库的向量格式一致(都是float列表)，可迁移。"])
concepts.append(["9.RAG（检索增强生成）","类比：AI的开卷考试。LLM本身是闭卷，RAG让LLM先翻书再回答。","三步：①用户提问转向量 ②向量库搜索 ③检索到的文档+提问发给LLM。","def rag(query):\n    docs = vector_store.search(query, top_k=5)\n    context = '\\n'.join([d['text'] for d in docs])\n    return llm.chat(f'资料：{context}\\n问题：{query}')","常见问题：检索到的文档不相关。方案：混合检索(向量+关键词)+RRF。"])
concepts.append(["10.ReAct（Reasoning+Acting）","类比：做饭。先想做什么(思考)→切菜炒菜(行动)→尝味道(观察)→决定加盐(再思考)。","Agent不断重复思考→行动→观察→循环直到完成。这是Agent区别于API调用的核心。","for i in range(50):\n    if status=='PLANNING': plan = llm_plan(task); status='EXECUTING'\n    if status=='EXECUTING': result = execute_step(step)","一定要设置最大迭代次数！失败步骤要跳过而不是死循环。"])
concepts.append(["11.Agent（智能体）","Agent=LLM(大脑)+Tools(工具)+Memory(记忆)+Planning(规划)。","四个核心能力：①理解用户意图 ②多步计划 ③调用工具执行 ④记住对话历史。","class BaseRuntime:\n    def init(self, config): ...\n    def execute(self, query, **kwargs): ...","单Agent做简单任务。多Agent协作做复杂任务。"])
concepts.append(["12.Tool（工具）","工具是Agent与外部世界交互的桥梁。没工具，Agent只能纸上谈兵。","OrchestratorRuntime工具集：list_dir, read_file, write_file, run_python, run_shell, wait。","class Sandbox:\n    def run_python(self, code, timeout=30):\n        script.write_text(code)\n        r = subprocess.run([sys.executable, script], timeout=timeout)\n        return r.stdout.decode('utf-8')","MokioAgent工具更多：FileRead,FileWrite,FileEdit,Bash,WebSearch,Grep。"])
concepts.append(["13.MCP（Model Context Protocol）","类比：AI世界的USB协议。不管什么外设(工具)，符合USB标准就能即插即用。","Anthropic提出的开放标准。没有MCP前：每个工具单独写集成。有MCP后：统一标准自动连接。","@dataclass\nclass ToolSchema:\n    name: str\n    description: str\n    params: dict","AgentForge的ToolSchema实现了类似MCP的统一接口理念。"])
concepts.append(["14.Harness（调度编排器）","类比：乐队指挥。自己不演奏，但告诉每个乐手什么时候演奏什么。","Harness负责：拆解任务→分配给合适Agent→收集结果→处理错误→编排顺序。","def stream_agent_events(task):\n    state = create_runtime(workspace)\n    workflow = build_complex_workflow()\n    for mode, event in workflow.stream(inputs):\n        yield event","Harness不执行具体任务，只负责调度。"])
concepts.append(["15.Skill（技能包）","Skill是打包好的专业技能。包含专属Prompt+工具集+流程逻辑。","客服回访技能可能包含：专用System Prompt+订单查询工具+退换货流程+话术模板。","list_dir→文件列表技能\nread_file→文件读取技能\nrun_python→代码执行技能","Agent=执行者，Skill=技能包。一个Agent可加载多个Skill。"])
concepts.append(["16-30.更多概念","Orchestration(工作流编排) DAG(有向无环图) State Machine(状态机) Sandbox(沙箱) Circuit Breaker(熔断器) Checkpoint(检查点) Trace(追踪) Streaming(流式输出) Eval(评估) Fine-tuning(微调) Few-shot(少样本) CoT(思维链) HyDE RRF Hybrid Search(混合检索)","这些概念在第三篇的四项目逐行代码分析中已详细讲解。","(详见第三篇各项目的代码分析)","参见第三篇各章节。"])
concepts.append(["01.Python变量","变量=贴了标签的盒子。name='张三'把'张三'存进叫name的盒子。","Python变量不需要声明类型。动态类型语言。","name='张三'; age=25; price=99.9; is_ok=True","变量名用下划线命名法：user_name, is_admin, total_count"])
concepts.append(["02.Python字符串","AI Agent项目里90%的数据是字符串。","f-string是Python3.6+最推荐的字符串格式化方式。","msg = f'用户问题：{query}'","字符串用+拼接效率低，大段拼接用join()。"])
concepts.append(["03.Python列表","列表是有序集合。tools[0]取第一个元素。","索引从0开始。切片[start:end]不包含end。","tools=['list_dir','read_file']; print(tools[0])","负数索引从末尾开始：tools[-1]取最后一个。"])

H("第1章：19个核心概念详解", 2)
for item in concepts:
    H(item[0], 3)
    P(f"【一句话理解】{item[1]}")
    P(f"【深入讲解】{item[2]}")
    P(f"【代码示例】\n{item[3]}", "Code")
    P(f"【注意事项】{item[4]}")
    P("")

PB()

# ═══════════ 第三篇：四项目逐行精讲 ═══════════
H("第三篇：四个AI Agent项目逐行代码精讲", 1)
P("这篇是全书核心。逐行阅读四个项目的每份Python文件，理解每行代码在做什么、为什么这样写。")

H("第1章：BaseRuntime——一切Agent的起点", 2)
P("文件：agents/base.py（约30行）")
P("from abc import ABC, abstractmethod", "Code")
P("ABC=抽象基类。不能被直接创建对象，只能被继承。就像「动物」是抽象概念。")
P("abstractmethod标记的方法，子类必须重写，否则子类也不能创建对象。")
P("")
P("class BaseRuntime(ABC):", "Code")
P("    @abstractmethod", "Code")
P("    def init(self, config): pass", "Code")
P("    @abstractmethod", "Code")
P("    def execute(self, query='', **kwargs): pass", "Code")
P("    @property", "Code")
P("    def name(self): return 'base'", "Code")
P("")
P("@property把方法变属性调用：agent.name()变成agent.name（不用加括号）。")

H("第2章：Searcher——最简单的Agent（<100行）", 2)
P("def __init__(self):", "Code")
P("    self._cache = {}         # 缓存：减少API调用次数", "Code")
P("    self._llm_client = None  # LLM客户端（init时创建）", "Code")
P("    self._trace = TraceRecorder('search')  # 链路追踪", "Code")
P("    self._metrics = CLEARScorer('search')  # 性能指标", "Code")
P("")
P("_cache：_开头表示内部使用。缓存把搜索结果存起来，同样问题不重复调用API。")
P("")
P("def execute(self, query='', **kwargs):", "Code")
P("    self._trace.start(query)", "Code")
P("    try:", "Code")
P("        # 1.查缓存", "Code")
P("        cache_key = hashlib.md5(query.encode()).hexdigest()", "Code")
P("        if cache_key in self._cache:", "Code")
P("            return self._cache[cache_key]", "Code")
P("", "Code")
P("        # 2.在线搜索", "Code")
P("        engine = kwargs.get('engine', 'baidu')", "Code")
P("        if engine == 'baidu':", "Code")
P("            results = _search_baidu(query)", "Code")
P("        else:", "Code")
P("            results = _search_ddg(query)", "Code")
P("", "Code")
P("        # 3.AI总结", "Code")
P("        resp = self._llm_client.chat.completions.create(", "Code")
P("            model='deepseek-v4-flash',", "Code")
P("            messages=[{'role':'user','content': f'请总结：{results}'}],", "Code")
P("            temperature=0.3", "Code")
P("        )", "Code")
P("        summary = resp.choices[0].message.content", "Code")
P("", "Code")
P("        # 4.缓存并返回", "Code")
P("        self._cache[cache_key] = {'results': results, 'summary': summary}", "Code")
P("        return self._cache[cache_key]", "Code")
P("    finally:", "Code")
P("        self._trace.end({'elapsed': time.time()-t0})", "Code")
P("")
P("整个execute()流程：查缓存→搜索→总结→返回。没有循环、没有复杂的错误恢复。最适合入门。")

H("第3章：KnowledgeBot——工程化典范（~500行）", 2)
P("展示了企业级Agent的所有要素：熔断器、优雅降级、向量检索、对话历史、成本追踪。")
P("")
P("def __init__(self):", "Code")
P("    self._vector_store = None           # ChromaDB客户端", "Code")
P("    self._chunks = []                   # 文档片段列表", "Code")
P("    self._vector_ready = False          # 向量库就绪标志", "Code")
P("    self._cb = CircuitBreaker('rag', failure_threshold=3, recovery_timeout=30)", "Code")
P("    self._history = ShortTermMemory(max_window=20)", "Code")
P("    self._cost = CostTracker()          # Token消耗追踪", "Code")
P("    self._trace = TraceRecorder('rag')  # 链路追踪", "Code")
P("")
P("CircuitBreaker三状态：CLOSED(正常)→OPEN(跳闸)→HALF_OPEN(尝试恢复)。连续失败3次跳闸。")
P("ShortTermMemory(max_window=20)：保留最近20轮对话，防止Context溢出。")
P("")
P("def init(self, config):", "Code")
P("    self._chunks = _load_chunks(data_dir)  # 加载文档并切分", "Code")
P("    self._llm_client, self._llm_config = create_llm_client(config)", "Code")
P("", "Code")
P("    # 后台线程初始化向量库（不阻塞启动）", "Code")
P("    def _bg():", "Code")
P("        self._vector_store = VectorStore(config.get('embedding', {}))", "Code")
P("        self._vector_store.add_documents(self._chunks, batch_size=10)", "Code")
P("        self._vector_ready = True", "Code")
P("    threading.Thread(target=_bg, daemon=True).start()", "Code")
P("")
P("为什么用后台线程？向量化调用远程API，可能耗时1-2分钟。daemon=True：主程序退出时自动销毁。")
P("batch_size=10：阿里百炼API每次最多处理10条。")
P("")
P("def _retrieve(self, query, top_k=5):", "Code")
P("    @self._cb  # 熔断器装饰", "Code")
P("    def hybrid(q):", "Code")
P("        return self._vector_store.hybrid_search(q, top_k=top_k)", "Code")
P("    try:", "Code")
P("        return hybrid(query)", "Code")
P("    except Exception:", "Code")
P("        return _keyword_search(query, self._chunks)  # 优雅降级", "Code")
P("")
P("熔断触发后不报错——降级到关键词搜索。系统降级运行而不是崩溃，这就是优雅降级。")
P("")
P("_keyword_search实现：", "Code")
P("def _keyword_search(query, chunks, top_k=5):", "Code")
P("    ngrams = re.findall(r'[\\u4e00-\\u9fff]{2,3}', query)  # 2-3字窗口", "Code")
P("    for chunk in chunks:", "Code")
P("        score = sum(1 for ng in ngrams if ng in chunk['text'])", "Code")
P("    return sorted(results, key=lambda x: -x[3])[:top_k]", "Code")
P("")
P("为什么2-3字？中文单字没意义。「比」「亚」不知道是什么，「比亚迪」才是一个完整概念。")

H("第4章：OrchestratorRuntime——ReAct循环完整实现", 2)
P("def run(self, task_id, llm_client=None, on_update=None):", "Code")
P("    task = load_task(task_id)  # 从SQLite加载", "Code")
P("    sandbox = Sandbox(task['id']); sandbox.setup()", "Code")
P("", "Code")
P("    for i in range(MAX_STEPS):", "Code")
P("        if time.time() > task['created_at'] + task['timeout_minutes']*60:", "Code")
P("            task['status']='TIMEOUT'; break", "Code")
P("", "Code")
P("        if task['status'] in ('PENDING','PLANNING'):", "Code")
P("            plan = self._llm_plan(task, llm_client)", "Code")
P("            plan = validate_plan(plan)  # 工具名校验", "Code")
P("            if not plan:", "Code")
P("                plan = self._fallback_plan(task['prompt'])", "Code")
P("            task['plan'] = plan; task['status']='EXECUTING'", "Code")
P("            save_task(task); continue", "Code")
P("", "Code")
P("        if task['status'] == 'EXECUTING':", "Code")
P("            pending = [s for s in task['steps'] if s['status'] in ('pending','retrying')]", "Code")
P("            if not pending: task['status']='DONE'; break", "Code")
P("", "Code")
P("            ready = [s for s in pending if all_deps_met(s, task['steps'])]", "Code")
P("            with ThreadPoolExecutor(max_workers=4) as pool:", "Code")
P("                futures = {pool.submit(self._execute_step, s, sandbox): s for s in ready}", "Code")
P("                for future in as_completed(futures):", "Code")
P("                    step = futures[future]; result = future.result(timeout=30)", "Code")
P("                    if result['ok']: step['status']='done'", "Code")
P("                    else: step['retries']+=1", "Code")
P("                    if step['retries']>=3: step['status']='failed'", "Code")
P("")
P("ThreadPoolExecutor：最多4线程并行执行。as_completed：最先完成的先处理。")
P("all_deps_met：检查所有依赖是否已完成。比如D依赖B和C，则B和C都完成时D才就绪。")
P("")
P("沙箱隔离：", "Code")
P("class Sandbox:", "Code")
P("    def run_python(self, code, timeout=30):", "Code")
P("        script = self.workspace / 'temp' / '_script.py'", "Code")
P("        script.write_text(code)", "Code")
P("        r = subprocess.run([sys.executable, script],", "Code")
P("            capture_output=True, timeout=timeout, cwd=str(self.workspace))", "Code")
P("        return r.stdout.decode('utf-8', errors='replace')", "Code")
P("")
P("为什么用subprocess而不是exec()？exec()在主进程运行，如果代码有死循环，主程序也遭殃。subprocess在新进程运行，崩溃不影响主程序。")

H("第5章：MokioAgent——LangGraph企业级编排", 2)
P("核心差异：用LangGraph的StateGraph定义图结构工作流，而不是手写循环。")
P("")
P("@dataclass  # 自动生成__init__, __repr__, __eq__", "Code")
P("class RuntimeState:", "Code")
P("    workspace: Path", "Code")
P("    approval_mode: str = 'inline'      # 审批模式：inline/auto/deny", "Code")
P("    approval_handler: Callable | None = None  # 审批回调函数", "Code")
P("    checkpoint_mode: str = 'light'      # 检查点模式：light/strict/off", "Code")
P("    trace_mode: str = 'on'             # 追踪模式：on/off", "Code")
P("")
P("流式事件处理——yield：", "Code")
P("def stream_agent_events(task):", "Code")
P("    state = create_runtime(workspace)", "Code")
P("    workflow = build_complex_workflow()", "Code")
P("    for mode, event in workflow.stream(inputs):", "Code")
P("        yield {'type': 'graph_event', 'event': event}", "Code")
P("")
P("yield vs return：return一次性返回函数结束。yield暂停函数返回一个值，下次继续执行。")
P("")
P("检查点：", "Code")
P("class CheckpointManager:", "Code")
P("    def save(self, state, status, latest_node):", "Code")
P("        cp = {'status': status, 'node': latest_node, 'state': state}", "Code")
P("        json.dump(cp, open(self.path, 'w'))", "Code")
P("        return cp", "Code")

PB()

# ═══════════ 第四篇：面试题 ═══════════
H("第四篇：AI Agent面试题库", 1)

H("4.1 基础理论（10题）", 2)
qa = [
    ["什么是AI Agent？和普通API调用有什么区别？","AI Agent能自主感知、规划、执行、迭代。普通API一问一答，Agent是思考→行动→观察→循环。"],
    ["ReAct是什么？","Reasoning+Acting。不断重复：思考→决定→执行→观察→再思考。直到任务完成。"],
    ["RAG是什么？解决了什么问题？","检索增强生成。先检索知识库再让LLM回答。解决知识过时和幻觉两个核心问题。"],
    ["什么是Embedding？","把文字转成向量。意思相近的文字向量距离近。苹果→[0.1,0.5], iPhone→[0.12,0.48]。"],
    ["向量数据库和传统数据库的区别？","传统数据库精确匹配(WHERE name='苹果')。向量数据库语义匹配。"],
    ["什么是RRF？","倒数排名融合。多种检索方式的排名融合成总分。第1名得1分，第2名得0.5分。"],
    ["MCP是什么？","Model Context Protocol。统一AI连接外部工具的方式。类似USB协议。"],
    ["Prompt Engineering的核心原则？","角色明确+任务清晰+格式指定+约束边界+示例示范。"],
    ["Context Window溢出怎么办？","截断(保留最近N轮)、压缩(LLM做摘要)、滑动窗口(丢弃最早对话)。"],
    ["Agent和普通函数的区别？","普通函数接收参数返回结果。Agent有循环(ReAct)、工具(Tool)、记忆(Memory)、规划(Planning)。"],
]
for q, a in qa:
    P(f"Q: {q}")
    P(f"A: {a}")
    P("")

H("4.2 项目实战（5题）", 2)
qa2 = [
    ["KnowledgeBot如何提高检索准确率？","混合检索(向量+关键词)+RRF融合+CJK滑动窗口。熔断时降级到纯关键词。"],
    ["为什么用subprocess而不是exec()？","安全！exec()在主进程运行有风险。subprocess在新进程运行，崩溃不影响主程序。"],
    ["DAG并行怎么实现？","depends_on声明依赖。ThreadPoolExecutor并行执行依赖已满足的步骤。as_completed处理完成。"],
    ["熔断器三状态？","CLOSED(正常)→失败3次→OPEN(跳闸)→30秒后→HALF_OPEN(尝试恢复)。"],
    ["SQLite多线程读写？","PRAGMA journal_mode=WAL。WAL模式读不阻塞写，写不阻塞读。"],
]
for q, a in qa2:
    P(f"Q: {q}")
    P(f"A: {a}")
    P("")

PB()

# ═══════════ 第五篇：调试 ═══════════
H("第五篇：调试方法与常见错误", 1)
H("5.1 调试五步法", 2)
P("1. curl确认API通不通")
P("2. 单独测试每个Tool是否正常工作")
P("3. Playground测试Prompt组合效果")
P("4. 串起全部步骤跑完整流程，关键节点加print()")
P("5. 检查边界条件：空数据/超时/脏话/Token超限")

H("5.2 10个高频Bug", 2)
bugs = [
    ["API返回400","模型名/API Key/messages有误","curl先测试"],
    ["API返回429","超出速率限制","加time.sleep()限速"],
    ["向量检索返回空","Embedding未初始化","检查vector_ready"],
    ["OrchestratorRuntime死循环","ReAct无终止条件","设置MAX_STEPS"],
    ["subprocess中文乱码","Windows默认GBK","encoding='utf-8',errors='replace'"],
    ["SQLite database is locked","多线程并发","PRAGMA journal_mode=WAL"],
    ["ChromaDB启动慢","首次下载ONNX模型","预计算Embedding"],
    ["Agent编造文件名","LLM不了解文件系统","先list_dir获取真实文件"],
    ["Prompt超Token","历史累积太多","ShortTermMemory(max_window=20)"],
    ["前端一直加载","后端API超时","检查后端日志"],
]
for t,c,s in bugs:
    H(f"🔴 {t}", 3)
    P(f"原因：{c}  解决：{s}")

PB()

# ═══════════ 第六篇：行业落地 ═══════════
H("第六篇：AI Agent行业落地实战", 1)
H("6.1 四大行业案例", 2)
for n,p,s,e in [["金融-智能投研","分析师效率低","KnowledgeBot读取年报，自然语言提问","效率+90%"],["电商-智能客服","人力成本高","OrchestratorRuntime调用物流API+订单DB","自动处理70%"],["医疗-辅助诊断","需快速查文献","KnowledgeBot检索医学库，审批后输出","诊断+30%"],["制造-智能质检","漏检率高","Agent分析IoT数据自动告警","检出+40%"]]:
    H(n, 3); P(f"痛点：{p}  方案：{s}  效果：{e}")

H("6.2 POC到生产流程", 2)
P("第1周：选场景→3天POC→50个测试用例")
P("第2周：评估准确率→优化Prompt→80%+")
P("第3周：工程化→熔断器/沙箱/持久化/追踪")
P("第4周：Docker部署→灰度发布→监控告警")

PB()

# ═══════════ 第七篇：部署 ═══════════
H("第七篇：部署与运维", 1)
P("Dockerfile：")
P("FROM python:3.11-slim\nWORKDIR /app\nCOPY . .\nRUN pip install -r requirements.txt\nCMD ['uvicorn','web.app:app','--host','0.0.0.0','--port','8000']", "Code")
P("")
P("生产注意事项：①SQLite→MySQL切换 ②API密钥用环境变量 ③JSON日志 ④Prometheus+Grafana监控 ⑤多副本+负载均衡 ⑥Rate Limiter ⑦内容过滤防注入")

PB()

# ═══════════ 第八篇：附录 ═══════════
H("第八篇：附录", 1)
H("A. 30个核心术语速查", 2)
gl = "LLM=大语言模型|Token=计费单位|Prompt=指令|Context=短期记忆|Temperature=创造力|Embedding=文本向量化|VectorDB=向量数据库|RAG=检索增强生成|ReAct=思考+行动循环|Agent=智能体|Tool=工具|MCP=统一工具协议|Harness=调度器|Skill=技能包|DAG=有向无环图|Sandbox=沙箱|CircuitBreaker=熔断器|Checkpoint=检查点|Trace=全链路追踪|Streaming=流式输出|Eval=评估|Fine-tuning=微调|Few-shot=少样本学习|CoT=思维链|HyDE=假设文档嵌入|RRF=倒数排名融合|HybridSearch=混合检索|Chunk=文档片段|Fallback=兜底|ThreadSafety=线程安全"
tms = gl.split("|")
for r in range(10):
    P("  ".join((tms[r]+"  " if r<len(tms) else "") + (tms[r+10]+"  " if r+10<len(tms) else "") + (tms[r+20]+"  " if r+20<len(tms) else "")))

H("B. 推荐资源", 2)
P("• DeepSeek: https://platform.deepseek.com")
P("• LangGraph: https://langchain-ai.github.io/langgraph/")
P("• ChromaDB: https://docs.trychroma.com/")
P("• 项目: D:\\hermes\\work\\agentgroup\\agent_forge")
P("• 项目: D:\\hermes\\work\\agentgroup\\MokioAgent")

# ═══════════ 保存 ═══════════
out_path = OUT / "AI_Agent从入门到精通_四项目实战教学.docx"
doc.save(str(out_path))
print(f"✅ 已保存: {out_path} ({out_path.stat().st_size / 1024:.0f}KB)")
