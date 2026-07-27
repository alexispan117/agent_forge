"""
生成大厂面试优化版简历 — 针对 AI Agent 实习岗位
布局：现代双栏式，重点突出技术深度与项目成果
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(r"D:\hermes\work\agentgroup\ai-knowledge")
OUT.mkdir(parents=True, exist_ok=True)
doc = Document()

# ── 页面设置 ──
section = doc.sections[0]
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)

# ── 颜色 ──
PRIMARY = RGBColor(0x4F, 0x46, 0xE5)    # Deep Indigo
DARK = RGBColor(0x1E, 0x29, 0x3B)       # Slate 800
MUTED = RGBColor(0x64, 0x74, 0x8B)      # Slate 500
ACCENT = RGBColor(0x05, 0x9E, 0x69)     # Emerald 600

# ── 样式 ──
style = doc.styles["Normal"]
style.font.name = "Inter"
style.font.size = Pt(10)
style.paragraph_format.line_spacing = 1.3
style.font.color.rgb = DARK

# ── 辅助函数 ──

def add_section_bar(text):
    """带左侧色块的章节标题"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    # 左侧色块
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    left = pBdr.makeelement(qn('w:left'), {
        qn('w:val'): 'single', qn('w:sz'): '24',
        qn('w:space'): '8', qn('w:color'): '4F46E5',
    })
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(f"  {text}")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = DARK
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    run = p.add_run("▸ ")
    run.font.color.rgb = PRIMARY
    run.font.size = Pt(9.5)
    if bold_prefix:
        r2 = p.add_run(bold_prefix)
        r2.font.bold = True; r2.font.size = Pt(9.5)
        r3 = p.add_run(text)
        r3.font.size = Pt(9.5)
    else:
        r2 = p.add_run(text)
        r2.font.size = Pt(9.5)
    return p

def add_meta(items):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(4)
    for i, item in enumerate(items):
        if i > 0:
            s = p.add_run("  ·  ")
            s.font.size = Pt(9); s.font.color.rgb = MUTED
        r = p.add_run(item)
        r.font.size = Pt(9); r.font.color.rgb = MUTED
    return p

def set_cell_shading(cell, color):
    """设置表格单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

# ═══════════ 正文 ═══════════

# ── 头部：姓名 + 标题行 ──
header_table = doc.add_table(rows=1, cols=2)
header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 左列：姓名 + 求职意向
left_cell = header_table.cell(0, 0)
left_cell.width = Cm(10)
p = left_cell.paragraphs[0]
run = p.add_run("Edward · 张明远")
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = DARK
p2 = left_cell.add_paragraph()
r2 = p2.add_run("AI Agent 开发实习生 · 2026 届")
r2.font.size = Pt(12)
r2.font.color.rgb = PRIMARY
r2.font.bold = True

# 右列：联系方式
right_cell = header_table.cell(0, 1)
right_cell.width = Cm(5)
for line in [
    "📧 zhangmingyuan@email.com",
    "📱 138-xxxx-xxxx",
    "📍 上海",
    "🔗 github.com/agentforge",
]:
    p = right_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(line)
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED

# 删除表格边框
for row in header_table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tcBorders.append(border)
        tcPr.append(tcBorders)

# ── 专业概述 ──
add_section_bar("专业概述")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.3)
r = p.add_run(
    "信息工程专业应届毕业生，专注 AI Agent 方向，独立完成企业级多智能体编排平台 AgentForge 全栈建设。"
    "项目对标金融/制造/政务企业需求，实现 Supervisor-Worker 四节点微服务、MCP/A2A 双协议互联、"
    "CLEAR 五维评估与全链路异步化，具备可向企业 CTO 路演的生产级工程标准。"
    "熟悉 LangGraph 状态机编排、RAG 混合检索、FastAPI 异步后端、Docker 容器化部署。"
)
r.font.size = Pt(9.5)

# ── 技术能力 ──
add_section_bar("技术能力")

# 三列表格
skill_table = doc.add_table(rows=2, cols=3)
skill_table.alignment = WD_TABLE_ALIGNMENT.CENTER

skill_data = [
    ["AI / LLM", "架构与工程", "工具与平台"],
    [
        "• Prompt Engineering (CoT/ReAct)\n• RAG + 混合检索 (RRF)\n• LangGraph 状态机编排\n• MCP / A2A 协议实现\n• CLEAR 评估体系",
        "• Supervisor-Worker 微服务\n• FastAPI + asyncio 异步\n• SQLite + ChromaDB 持久化\n• Docker-Compose 容器编排\n• CI/CD · GitHub Actions",
        "• Python / C++ / TypeScript\n• DeepSeek / OpenAI API\n• 阿里 DashScope Embedding\n• ECharts + SSE 实时推送\n• LangChain / LangGraph"
    ],
]

for i, row in enumerate(skill_table.rows):
    for j, cell in enumerate(row.cells):
        cell.text = skill_data[i][j] if j < len(skill_data[i]) else ""
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)
            for run in paragraph.runs:
                run.font.size = Pt(8.5)
                run.font.color.rgb = DARK
                if i == 0:
                    run.font.bold = True
                    run.font.color.rgb = PRIMARY
                    run.font.size = Pt(9.5)

# 设置列宽
for row in skill_table.rows:
    row.cells[0].width = Cm(5)
    row.cells[1].width = Cm(5)
    row.cells[2].width = Cm(5)

# 去除表格边框
for row in skill_table.rows:
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for bn in ['top','left','bottom','right','insideH','insideV']:
            b = OxmlElement(f'w:{bn}')
            b.set(qn('w:val'), 'nil')
            tcBorders.append(b)
        tcPr.append(tcBorders)

# ── 核心项目 ──
add_section_bar("核心项目")

# AgentForge
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
r = p.add_run("AgentForge — 企业级多智能体编排平台")
r.font.bold = True
r.font.size = Pt(11)
r.font.color.rgb = DARK

add_meta(["2026.03 – 2026.07", "独立架构设计 + 全栈开发", "项目地址: github.com/agentforge"])

bullets_af = [
    ("微服务架构: ",
     "Supervisor + 3 Worker 四容器编排（Docker-Compose），代码卷挂载支持热重载。"
     "自研 MCP (JSON-RPC) + A2A (AgentCard) 双协议层，Agent 间零硬编码互联，支持横向扩展与故障隔离。"),
    ("韧性工程: ",
     "@with_fallback 三级防线（重试→熔断→本地降级），演示中杀死 Worker 容器后工作流仍 100% 交付。"
     "所有 I/O 调用标注 # ASYNC 注释并锁定库版本，实现可审计的异步化改造。"),
    ("评估体系: ",
     "CLEAR 五维评估（成本/延迟/效能/保证/可靠性），全部由真实执行数据计算，ECharts 雷达图实时渲染。"
     "21 个 pytest 用例 + GitHub Actions CI 保障工程质量。密钥全量外移至环境变量。"),
    ("演示体验: ",
     "离线 Mock 模式（LLM_MOCK_MODE=true）默认开启，7 个预置响应模板覆盖全流程，零网络即可运行。"
     "前端 Jinja2 + SSE 实时推送 + ECharts 白底 SaaS 仪表盘，Inter + JetBrains Mono 字体本地加载。"),
]
for bold, text in bullets_af:
    add_bullet(text, bold)

# 量化成果
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.6)
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(2)
r = p.add_run("📊 量化成果: ")
r.font.bold = True; r.font.size = Pt(9.5)
items = [
    "复杂多步任务完成率 50% → 75%（21 个自动化用例）",
    "RAG 增强后幻觉率降低 ~15%",
    "四容器联调全链路 200 通过",
    "CLEAR 评分如实反映故障注入场景",
]
for i, item in enumerate(items):
    if i > 0:
        p.add_run("  |  ").font.size = Pt(9)
    r2 = p.add_run(item)
    r2.font.size = Pt(9); r2.font.color.rgb = ACCENT

# 智能小车
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
r = p.add_run("智能小车自主决策系统")
r.font.bold = True; r.font.size = Pt(10.5)
add_meta(["2024.06 – 2024.12", "嵌入式 C++", "独立开发"])
add_bullet("多传感器融合的「感知→决策→执行」闭环，状态机实现避障/循迹/路径规划动态切换。"
           "架构与 Agent 的 Planning–Tool Use–Execution 范式同源，奠定资源受限环境下的可靠决策系统思维。")

# 数据库项目
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
r = p.add_run("数据库系统设计与优化")
r.font.bold = True; r.font.size = Pt(10.5)
add_meta(["2025.09 – 2025.10", "课程项目"])
add_bullet("SQL Server 架构设计、索引优化与客户端开发。为后续 ChromaDB 向量检索与混合检索调优建立工程直觉。")

# ── 技术影响力 ──
add_section_bar("技术影响力")
add_bullet("开源项目: AgentForge 完整企业级多 Agent 编排平台，可供面试官现场审查源码")
add_bullet("技术布道: 独立编写 20,000+ 字 AI Agent 教学文档，覆盖四项目逐行代码讲解 + 面试题库 + 行业案例")
add_bullet("前沿追踪: MCP/A2A 协议（已落地）| LangGraph | Harness Engineering | Reflexion 反思机制 | AutoGen/CrewAI")

# ── 教育背景 ──
add_section_bar("教育背景")
p = doc.add_paragraph()
r = p.add_run("上海大学")
r.font.bold = True; r.font.size = Pt(11)
rs = p.add_run("  ·  信息工程专业（本科）  ·  2022.09 – 2026.06")
rs.font.size = Pt(10.5)

add_bullet("主修课程: 数据结构与算法、数据库系统、机器学习基础、嵌入式系统设计")
add_bullet("研究方向: 大语言模型应用开发与多智能体架构设计")
add_bullet("英语: CET-6，熟练阅读英文技术文档与学术论文")

# ── 实习经历 ──
add_section_bar("实习经历")
p = doc.add_paragraph()
r = p.add_run("金工实习 — 智能制造流程实践")
r.font.bold = True; r.font.size = Pt(10.5)
add_meta(["上海大学工程训练中心", "2023.06 – 2023.08"])
add_bullet("参与数控机床、3D 打印、CATIA 设计的完整工业制造流程，理解制造业现场痛点。"
           "为「AI + 工业制造」场景落地提供一线认知，培养软硬件协同思维。")

# ── 保存 ──
out_path = OUT / "Edward_Resume_AI_Agent_大厂优化版.docx"
doc.save(str(out_path))
print(f"✅ 简历已生成: {out_path}")
print(f"   大小: {out_path.stat().st_size // 1024}KB")
