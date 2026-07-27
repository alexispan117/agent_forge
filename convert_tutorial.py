"""
将 AgentForge_Complete_Tutorial.md 转换为格式化的 .docx
处理：标题、代码块、表格、列表、引用、粗斜体
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

MD_PATH = Path(r"D:\hermes\work\agentgroup\AgentForge_Complete_Tutorial.md")
OUT_DIR = Path(r"D:\hermes\work\agentgroup")
OUT_PATH = OUT_DIR / "AgentForge_Complete_Tutorial.docx"

doc = Document()

# ── 样式 ──
style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.4

for lv in range(1, 5):
    s = doc.styles[f"Heading {lv}"]
    s.font.name = "微软雅黑"

cs = doc.styles.add_style("Code", 1)
cs.font.name = "Consolas"
cs.font.size = Pt(8.5)
cs.paragraph_format.space_before = Pt(1)
cs.paragraph_format.space_after = Pt(1)

# ── 读取并逐行处理 ──
lines = MD_PATH.read_text(encoding="utf-8").splitlines()

in_code_block = False
code_buffer = []
table_buffer = []
in_table = False
table_col_count = 0

def flush_code():
    if code_buffer:
        text = "\n".join(code_buffer)
        doc.add_paragraph(text, style="Code")
        code_buffer.clear()

def flush_table():
    if not table_buffer:
        return
    if len(table_buffer) < 2:
        table_buffer.clear()
        return
    # Determine column count from header row
    header_cells = [c.strip() for c in table_buffer[0].split("|") if c.strip()]
    cols = len(header_cells)
    if cols < 2:
        table_buffer.clear()
        return
    
    rows_data = []
    for line in table_buffer:
        cells = [c.strip() for c in line.split("|") if c.strip()]
        if cells:
            rows_data.append(cells)
    
    if len(rows_data) < 2:
        table_buffer.clear()
        return
    
    # Trim rows to same column count
    max_cols = max(len(r) for r in rows_data)
    rows_data = [r + [""] * (max_cols - len(r)) for r in rows_data]
    
    t = doc.add_table(rows=len(rows_data), cols=max_cols)
    t.style = 'Table Grid'
    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            if j < len(t.rows[i].cells):
                cell = t.rows[i].cells[j]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.name = "微软雅黑"
                        if i == 0:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0x63, 0x66, 0xF1)
    doc.add_paragraph()  # spacing after table
    table_buffer.clear()

line_count = len(lines)
last_progress = 0

for idx, line in enumerate(lines):
    # Progress
    progress = (idx + 1) * 100 // line_count
    if progress >= last_progress + 5:
        last_progress = progress
        print(f"\r  处理中: {progress}% ({idx+1}/{line_count})", end="", flush=True)

    stripped = line.strip()
    
    # ── 代码块切换 ──
    if stripped.startswith("```"):
        if in_code_block:
            flush_code()
            in_code_block = False
        else:
            flush_table()
            in_code_block = True
        continue
    
    if in_code_block:
        code_buffer.append(line)
        continue
    
    # ── 表格 ──
    if "|" in stripped and stripped.count("|") >= 3 and stripped.startswith("|"):
        # Check if it's a separator row (|---|)
        if re.match(r'^\|[\s\-:]+\|', stripped) and "---" in stripped:
            continue  # skip separator rows
        table_buffer.append(stripped)
        in_table = True
        continue
    else:
        if in_table:
            flush_table()
            in_table = False
    
    # ── 引用 ──
    if stripped.startswith(">"):
        text = stripped.lstrip("> ").strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        run.font.italic = True
        continue
    
    # ── 空行 ──
    if not stripped:
        continue
    
    # ── 标题 ──
    if stripped.startswith("## "):
        flush_code()
        h = doc.add_heading(stripped[3:], 2)
        continue
    if stripped.startswith("### "):
        flush_code()
        h = doc.add_heading(stripped[4:], 3)
        continue
    if stripped.startswith("#### "):
        flush_code()
        h = doc.add_heading(stripped[5:], 4)
        continue
    if stripped.startswith("# "):
        flush_code()
        h = doc.add_heading(stripped[2:], 1)
        continue
    
    # ── 列表 ──
    if stripped.startswith("- ") or stripped.startswith("* "):
        text = stripped[2:]
        # Remove bold markers for cleaner text
        text = text.replace("**", "")
        p = doc.add_paragraph(text, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(10.5)
        continue
    
    if stripped.startswith("1. ") or stripped.startswith("2. ") or stripped.startswith("3. ") or re.match(r'^\d+\.\s', stripped):
        text = re.sub(r'^\d+\.\s', '', stripped)
        text = text.replace("**", "")
        p = doc.add_paragraph(text, style="List Number")
        for run in p.runs:
            run.font.size = Pt(10.5)
        continue
    
    # ── 分割线 ──
    if stripped in ("---", "***", "___"):
        flush_code()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        # Add a bottom border to simulate a horizontal rule
        pPr = p._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn('w:pBdr'), {})
        bottom = pBdr.makeelement(qn('w:bottom'), {
            qn('w:val'): 'single',
            qn('w:sz'): '6',
            qn('w:space'): '1',
            qn('w:color'): 'E2E8F0',
        })
        pBdr.append(bottom)
        pPr.append(pBdr)
        continue
    
    # ── 普通段落 ──
    text = line
    # Handle inline code
    text = text.replace("**", "")
    text = text.replace("`", "")
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.name = "微软雅黑"

# Flush remaining
flush_code()
flush_table()

print(f"\r  处理中: 100% ({line_count}/{line_count})")

# ── 保存 ──
doc.save(str(OUT_PATH))
print(f"\n✅ 转换完成: {OUT_PATH}")
print(f"   源文件: {line_count} 行, {MD_PATH.stat().st_size // 1024}KB")
print(f"   DOCX: {OUT_PATH.stat().st_size // 1024}KB")
